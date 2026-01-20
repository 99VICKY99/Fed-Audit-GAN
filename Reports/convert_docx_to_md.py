"""Convert all .docx files in the Reports folder to Markdown format."""
from docx import Document
import os
from pathlib import Path

def docx_to_markdown(docx_path):
    """Convert a .docx file to Markdown format."""
    doc = Document(docx_path)
    markdown_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            markdown_content.append("")
            continue
            
        # Handle headings
        if para.style.name.startswith('Heading'):
            level = para.style.name.replace('Heading ', '')
            try:
                level_num = int(level)
                markdown_content.append(f"{'#' * level_num} {text}")
            except:
                markdown_content.append(f"## {text}")
        else:
            markdown_content.append(text)
    
    # Handle tables
    for table in doc.tables:
        markdown_content.append("\n")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            markdown_content.append("| " + " | ".join(cells) + " |")
            if i == 0:  # Add separator after header
                markdown_content.append("| " + " | ".join(["---"] * len(cells)) + " |")
        markdown_content.append("\n")
    
    return "\n".join(markdown_content)

# Convert all .docx files in the current directory
reports_dir = Path(__file__).parent
for docx_file in reports_dir.glob("*.docx"):
    print(f"Converting {docx_file.name}...")
    try:
        md_content = docx_to_markdown(docx_file)
        md_file = docx_file.with_suffix('.md')
        md_file.write_text(md_content, encoding='utf-8')
        print(f"  ✓ Created {md_file.name}")
    except Exception as e:
        print(f"  ✗ Error: {e}")

print("\nConversion complete!")
